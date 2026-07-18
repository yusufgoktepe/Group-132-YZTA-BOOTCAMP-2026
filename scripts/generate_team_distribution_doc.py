from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT_DIR = Path("artifacts")
DOCX_PATH = OUTPUT_DIR / "ekip-gorev-dagilimi.pdf-source.docx"


NAVY = RGBColor(16, 43, 73)
SLATE = RGBColor(82, 96, 109)
BLUE = RGBColor(46, 116, 181)
LIGHT = RGBColor(240, 244, 248)
MID = RGBColor(217, 225, 232)
DARK = RGBColor(45, 45, 45)
WHITE = RGBColor(255, 255, 255)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for key, value in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tc_mar.find(qn(f"w:{key}"))
        if node is None:
            node = OxmlElement(f"w:{key}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color="DADCE0", size="8"):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:color"), color)


def add_run(paragraph, text, *, bold=False, size=11, color=DARK, font="Calibri", italic=False):
    run = paragraph.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.name = font
    run._element.rPr.rFonts.set(qn("w:ascii"), font)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), font)
    run.font.size = Pt(size)
    run.font.color.rgb = color
    return run


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Inches(0.25 + level * 0.2)
    p.paragraph_format.first_line_indent = Inches(-0.18)
    add_run(p, text, size=11)


def add_number(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.first_line_indent = Inches(-0.18)
    add_run(p, text, size=11)


def style_document(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.49)
    section.footer_distance = Inches(0.49)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = DARK
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15

    for name, size, color in [
        ("Heading 1", 16, BLUE),
        ("Heading 2", 13, BLUE),
        ("Heading 3", 12, NAVY),
    ]:
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(12)
        style.paragraph_format.space_after = Pt(6)


def add_header_footer(doc):
    section = doc.sections[0]
    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_run(header, "CampusMatch AI | Ekip Planı", size=9, color=SLATE)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(footer, "Group 132 - YZTA Bootcamp 2026", size=9, color=SLATE)


def add_title_block(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(4)
    add_run(p, "CampusMatch AI", size=24, bold=True, color=NAVY)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(14)
    add_run(p, "Ekip Görev Dağılımı ve Çalışma Planı", size=14, color=SLATE)

    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Inches(6.5)
    set_table_borders(table, color="FFFFFF", size="0")
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F0F4F8")
    set_cell_margins(cell, 120, 160, 120, 160)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(0)
    add_run(
        p,
        "Amaç: Projeyi ana hatlara bölmek, 5 kişilik ekip içinde sorumlulukları netleştirmek ve Sprint 2 sürecinde paralel ama kontrollü bir ilerleme sağlamak.",
        size=11,
        color=NAVY,
    )


def add_workstreams(doc):
    doc.add_paragraph("Projenin Ana Hatları", style="Heading 1")
    for item in [
        "Ürün ve proje yönetimi",
        "Mobil uygulama geliştirme",
        "Backend ve API geliştirme",
        "AI / recommendation ve veri tarafı",
        "Dokümantasyon, test ve demo hazırlığı",
    ]:
        add_bullet(doc, item)


def add_team_table(doc):
    doc.add_paragraph("Önerilen Ekip Dağılımı", style="Heading 1")
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    widths = [Inches(1.1), Inches(2.1), Inches(3.3)]
    for i, width in enumerate(widths):
        table.columns[i].width = width
    set_table_borders(table)

    headers = ["Kişi", "Ana Sorumluluk", "Temel Odak"]
    for i, text in enumerate(headers):
        cell = table.cell(0, i)
        set_cell_shading(cell, "102B49")
        set_cell_margins(cell)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_run(p, text, bold=True, color=WHITE)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    rows = [
        ("Kişi 1", "Proje yönetimi / Product", "Backlog, sprint planı, user flow, demo senaryosu"),
        ("Kişi 2", "Mobil geliştirme 1", "Expo yapısı, navigation, onboarding, profil ekranı"),
        ("Kişi 3", "Mobil geliştirme 2 + entegrasyon", "Öneriler, detay, kaydedilenler, backend bağlantısı"),
        ("Kişi 4", "Backend / API", "FastAPI, endpointler, veri modelleri, response yapıları"),
        ("Kişi 5", "AI / veri + dokümantasyon desteği", "Recommendation scoring, sentetik veri, test verileri, README desteği"),
    ]
    for idx, row in enumerate(rows):
        cells = table.add_row().cells
        for i, text in enumerate(row):
            set_cell_margins(cells[i])
            if idx % 2 == 1:
                set_cell_shading(cells[i], "F8FAFC")
            p = cells[i].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            add_run(p, text, size=10.5)
            cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def add_role_sections(doc):
    doc.add_paragraph("Rol Bazlı Sorumluluklar", style="Heading 1")
    sections = [
        ("1. Proje Yönetimi / Product", [
            "Backlog yönetimi ve görev dağılımı",
            "Sprint planlama ve öncelik sıralama",
            "User flow ve MVP kapsamını güncel tutma",
            "Demo senaryosu ve ekip koordinasyonu",
        ]),
        ("2. Mobil Geliştirme 1", [
            "React Native / Expo temel yapısı",
            "Navigation kurulumu",
            "Onboarding ve profil oluşturma ekranı",
            "Ortak component yapısı",
        ]),
        ("3. Mobil Geliştirme 2 + Entegrasyon", [
            "Etkinlik önerileri ekranı",
            "Etkinlik detay ekranı ve kaydedilenler",
            "Backend entegrasyonu",
            "UI düzenlemeleri ve kullanıcı akışı iyileştirmeleri",
        ]),
        ("4. Backend / API", [
            "FastAPI proje yapısı",
            "GET /health, GET /students, GET /clubs, GET /events",
            "POST /recommendations/student/{student_id}",
            "Veri modelleri ve response yapıları",
        ]),
        ("5. AI / Veri + Dokümantasyon", [
            "Sentetik veri yapısı ve örnek test verileri",
            "Recommendation scoring mantığı",
            "Açıklanabilir öneri metni",
            "README ve sprint doküman desteği",
        ]),
    ]
    for title, items in sections:
        doc.add_paragraph(title, style="Heading 2")
        for item in items:
            add_bullet(doc, item)


def add_working_principles(doc):
    doc.add_paragraph("Çalışma Prensibi ve Öncelik Sırası", style="Heading 1")
    for item in [
        "Herkesin bir ana sorumluluğu olacak.",
        "Her işin tek sahibi olacak, gerekirse destekli yürütülecek.",
        "Bağımlı işler yüzünden bekleme olursa mock data veya fake response ile ilerleme yapılacak.",
        "Öncelik öğrenci tarafı MVP olacak; kulüp yöneticisi tarafı ikinci aşamada genişletilecek.",
    ]:
        add_bullet(doc, item)

    doc.add_paragraph("Öncelik Sırası", style="Heading 2")
    for item in [
        "Ürün netliği ve backlog",
        "Öğrenci tarafı mobil ekranlar",
        "Backend endpointleri",
        "Recommendation mantığı",
        "Entegrasyon",
        "Demo ve dokümantasyon",
    ]:
        add_number(doc, item)


def add_status_logic(doc):
    doc.add_paragraph("Kısa Durum Mantığı", style="Heading 1")
    for item in [
        "Mobil ekip, backend tamamen bitmeden mock veriyle çalışabilir.",
        "Backend ekip, AI tamamen bitmeden fake response ile ilerleyebilir.",
        "AI ekip, mobil tamamen hazır olmadan veri ve scoring testleri yapabilir.",
        "Product tarafı tüm bu süreci birbirine bağlar.",
    ]:
        add_bullet(doc, item)

    doc.add_paragraph("Haftalık Mini Çalışma Şablonu", style="Heading 2")
    for item in [
        "Haftanın başı: görev dağılımı ve öncelik netleştirme",
        "Hafta ortası: kısa ilerleme kontrolü",
        "Hafta sonu: çıkan işlerin birleştirilmesi ve eksiklerin görülmesi",
    ]:
        add_number(doc, item)


def main():
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document()
    style_document(doc)
    add_header_footer(doc)
    add_title_block(doc)
    add_workstreams(doc)
    add_team_table(doc)
    add_role_sections(doc)
    add_working_principles(doc)
    add_status_logic(doc)
    doc.save(DOCX_PATH)
    print(DOCX_PATH)


if __name__ == "__main__":
    main()
